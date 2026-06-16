from __future__ import annotations

import io
import json
import os
import re
import time
from datetime import date
from collections import OrderedDict
from html import escape as html_escape
from typing import Dict, Iterable, List, Optional, Set

import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt
from pypdf import PdfReader

try:
    from agents.openai_client import get_agent
except Exception:  # pragma: no cover - dependency/setup guard
    def get_agent(*args, **kwargs):
        raise RuntimeError("OpenAI/LangChain dependencies are not installed. Run `pip install -r requirements.txt`.")

try:
    from langchain_community.document_loaders import AsyncChromiumLoader
    from langchain_community.document_transformers import Html2TextTransformer
except Exception:  # pragma: no cover - optional URL extraction dependency
    AsyncChromiumLoader = None
    Html2TextTransformer = None

try:
    from prompts import job_listing_prompt
except Exception:  # pragma: no cover - optional LLM cleanup dependency
    job_listing_prompt = None
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

try:
    import trafilatura
except Exception:  # pragma: no cover - optional fallback
    trafilatura = None

load_dotenv()
os.environ.setdefault("USER_AGENT", "AI-Cover-Letter-Generator/1.0 (+https://github.com/aakanksha105/Cover-Letter-Generator)")


SKILL_CATALOG: Dict[str, List[str]] = OrderedDict(
    {
        "Programming": [
            "Python", "SQL", "JavaScript", "TypeScript", "Java", "C++", "C", "C#", "Go", "R", "Bash", "PHP", "Ruby", "Scala",
        ],
        "Backend & APIs": [
            "REST API", "API", "Flask", "FastAPI", "Django", "Node.js", "Express.js", "Spring Boot",
            "Microservices", "OpenAPI", "Swagger", "JWT", "GraphQL", "Postman", "Serverless", "Webhooks",
        ],
        "Frontend": [
            "React", "React.js", "Next.js", "HTML", "CSS", "Tailwind CSS", "Bootstrap", "Material UI", "Streamlit", "Angular", "Vue.js",
        ],
        "Databases": [
            "PostgreSQL", "MySQL", "MongoDB", "SQLite", "Neo4j", "BigQuery", "Redis", "DynamoDB", "Snowflake", "Oracle",
        ],
        "Data & BI": [
            "Pandas", "NumPy", "Power BI", "Tableau", "Excel", "Data Analysis", "EDA", "ETL", "ELT", "Data Pipeline",
            "Data Pipelines", "Dashboard", "Dashboards", "Reporting", "Analytics", "Data Visualization", "Data Modeling", "A/B Testing",
        ],
        "AI & ML": [
            "Machine Learning", "Scikit-learn", "TensorFlow", "Keras", "PyTorch", "NLP", "BERT", "LLM", "OpenAI", "LangChain",
            "Generative AI", "Prompt Engineering", "Transformers", "RAG", "Vector Database", "Embeddings",
        ],
        "Cloud & DevOps": [
            "Docker", "Docker Compose", "AWS", "Azure", "GCP", "Git", "GitHub", "CI/CD", "GitHub Actions", "Linux",
            "Kubernetes", "Terraform", "Jenkins", "Azure DevOps", "CloudFormation",
        ],
        "Testing & Practices": [
            "Unit Testing", "Integration Testing", "Testing", "Pytest", "JUnit", "Jest", "Agile", "Scrum", "Code Review", "Debugging", "Documentation",
        ],
    }
)

ALIASES: Dict[str, List[str]] = {
    "REST API": ["rest api", "rest apis", "restful api", "restful apis", "api endpoint", "api endpoints"],
    "API": ["api", "apis"],
    "APIs": ["api", "apis"],
    "Node.js": ["node.js", "nodejs", "node js", "node runtime"],
    "Express.js": ["express.js", "expressjs", "express"],
    "React.js": ["react.js", "reactjs", "react"],
    "Next.js": ["next.js", "nextjs", "next js"],
    "Vue.js": ["vue.js", "vuejs", "vue"],
    "Tailwind CSS": ["tailwind", "tailwindcss", "tailwind css"],
    "PostgreSQL": ["postgresql", "postgres"],
    "Power BI": ["power bi", "powerbi"],
    "Scikit-learn": ["scikit-learn", "sklearn", "scikit learn"],
    "CI/CD": ["ci/cd", "cicd", "ci cd", "continuous integration", "continuous deployment"],
    "GitHub Actions": ["github actions", "gh actions"],
    "Machine Learning": ["machine learning", "ml"],
    "Generative AI": ["generative ai", "gen ai", "genai"],
    "Data Pipeline": ["data pipeline", "data pipelines", "etl pipeline", "etl pipelines"],
    "Dashboard": ["dashboard", "dashboards", "visualization", "visualizations"],
    "Unit Testing": ["unit testing", "unit tests", "pytest", "junit", "jest"],
    "Integration Testing": ["integration testing", "integration tests"],
    "Vector Database": ["vector database", "vector db", "pinecone", "weaviate", "chroma"],
}

_CANONICAL_BY_ALIAS: Dict[str, str] = {}
for _category, _skills in SKILL_CATALOG.items():
    for _skill in _skills:
        _CANONICAL_BY_ALIAS[_skill.lower()] = _skill
        for _alias in ALIASES.get(_skill, []):
            _CANONICAL_BY_ALIAS[_alias.lower()] = _skill

GENERIC_STOPWORDS = {
    "team", "teams", "work", "working", "role", "roles", "job", "jobs", "candidate", "candidates", "experience", "years",
    "strong", "excellent", "good", "ability", "skills", "skill", "required", "preferred", "responsibilities", "requirements",
    "benefits", "office", "remote", "hybrid", "salary", "communication", "collaboration", "problem", "solving", "fast", "paced",
    "environment", "employee", "employees", "customer", "customers", "business", "product", "products", "service", "services", "including",
    "support", "create", "build", "develop", "design", "manage", "using", "based", "knowledge", "familiarity", "understanding",
}

JOB_MARKERS = [
    "responsibilities", "requirements", "qualifications", "preferred qualifications", "what you'll do", "what you will do",
    "about the role", "about this role", "skills", "experience", "minimum qualifications", "basic qualifications",
]


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _strip_html(text: str) -> str:
    text = text or ""
    # Avoid BeautifulSoup warnings when the value is plain text or a URL/path-like string.
    if "<" not in text and ">" not in text:
        return text
    return BeautifulSoup(text, "html.parser").get_text("\n")


def _read_docx(file) -> str:
    doc = Document(file)
    blocks: List[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    # Many resume templates place the entire resume inside DOCX tables.
    # python-docx does not include table cell text in doc.paragraphs, so read it explicitly.
    for table in doc.tables:
        for row in table.rows:
            cell_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
                if cell.text and cell.text.strip()
            )
            if cell_text:
                blocks.append(cell_text)

    text = "\n".join(blocks)
    if not text.strip():
        raise ValueError("No readable text was found in the uploaded DOCX.")
    return text.strip()




def assess_resume_extraction_quality(text: str) -> Dict[str, object]:
    """Return simple quality signals so the UI can warn users when parsing looks unreliable."""
    text = text or ""
    words = re.findall(r"[A-Za-z0-9+#.]+", text)
    word_count = len(words)
    strange_chars = len(re.findall(r"[^\w\s.,;:!?@/#&()'\-+|]", text))
    strange_ratio = strange_chars / max(len(text), 1)
    lower = text.lower()
    sections_found = [
        label
        for label in ["experience", "skills", "education", "projects", "work", "employment"]
        if label in lower
    ]
    warnings: List[str] = []
    if word_count < 250:
        warnings.append("Extracted resume text is short; important details may be missing.")
    if strange_ratio > 0.08:
        warnings.append("Extracted resume text has unusual characters; the resume layout may not have parsed cleanly.")
    if len(sections_found) < 2:
        warnings.append("Few standard resume sections were detected.")
    return {
        "word_count": word_count,
        "strange_character_ratio": round(strange_ratio, 3),
        "sections_found": sections_found,
        "warnings": warnings,
        "is_likely_good": not warnings,
    }


def normalize_pasted_resume_text(text: str) -> str:
    """Clean manually pasted resume text while preserving useful line breaks."""
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_resume(file) -> str:
    name = getattr(file, "name", "").lower()
    if name.endswith(".docx"):
        return _read_docx(file)

    reader = PdfReader(file)
    pages: List[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text)

    resume_text = "\n".join(pages).strip()
    if not resume_text:
        raise ValueError("No readable text was found in the uploaded PDF.")
    return resume_text


def _extract_json_ld_jobposting(html: str) -> str:
    soup = BeautifulSoup(html or "", "html.parser")
    blocks: List[str] = []
    for script in soup.find_all("script", type="application/ld+json"):
        content = script.string or script.get_text(" ", strip=True)
        if not content:
            continue
        try:
            data = json.loads(content)
        except Exception:
            continue

        queue = data if isinstance(data, list) else [data]
        expanded = []
        for item in queue:
            if isinstance(item, dict) and isinstance(item.get("@graph"), list):
                expanded.extend(item["@graph"])
            expanded.append(item)

        for item in expanded:
            if not isinstance(item, dict):
                continue
            type_value = item.get("@type", "")
            type_text = " ".join(type_value) if isinstance(type_value, list) else str(type_value)
            if "JobPosting" not in type_text:
                continue
            fields = []
            title = item.get("title")
            if title:
                fields.append(f"Job title: {_strip_html(str(title))}")
            org = item.get("hiringOrganization")
            if isinstance(org, dict) and org.get("name"):
                fields.append(f"Company: {_strip_html(str(org['name']))}")
            for key in ["description", "responsibilities", "skills", "qualifications", "experienceRequirements", "educationRequirements"]:
                val = item.get(key)
                if isinstance(val, dict):
                    val = "\n".join(str(v) for v in val.values())
                elif isinstance(val, list):
                    val = "\n".join(str(v) for v in val)
                if val:
                    fields.append(_strip_html(str(val)))
            if fields:
                blocks.append("\n".join(fields))
    return "\n\n".join(blocks)


def _clean_webpage_text(text: str) -> str:
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text or "")
    text = re.sub(r"(?i)(accept cookies|cookie settings|privacy policy|terms of use|sign in|log in|create alert|subscribe|share this job|save job|apply now|equal opportunity employer|eeo|accessibility)", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    lines = []
    seen = set()
    for line in text.splitlines():
        cleaned = line.strip(" \t•·|-–—")
        if len(cleaned) < 3:
            continue
        low = cleaned.lower()
        if low in seen:
            continue
        if len(cleaned.split()) <= 2 and not any(marker in low for marker in JOB_MARKERS):
            continue
        seen.add(low)
        lines.append(cleaned)
    return "\n".join(lines).strip()


def _job_text_score(text: str) -> int:
    low = (text or "").lower()
    score = min(len(text or "") // 120, 100)
    score += sum(18 for marker in JOB_MARKERS if marker in low)
    score += sum(10 for tech in ["python", "sql", "api", "aws", "react", "java", "data", "dashboard", "machine learning", "docker"] if tech in low)
    score -= sum(16 for noise in ["cookie", "privacy", "subscribe", "newsletter", "recommended jobs", "similar jobs"] if noise in low)
    return score


def _requests_fallback(url: str) -> str:
    headers = {"User-Agent": os.environ.get("USER_AGENT", "AI-Cover-Letter-Generator/1.0")}
    response = requests.get(url, timeout=20, headers=headers)
    response.raise_for_status()
    html = response.text
    candidates: List[str] = []

    json_ld = _extract_json_ld_jobposting(html)
    if len(json_ld) > 300:
        candidates.append(json_ld)

    if trafilatura is not None:
        extracted = trafilatura.extract(html, include_comments=False, include_tables=False, favor_recall=True)
        if extracted:
            candidates.append(extracted)

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript", "svg"]):
        tag.decompose()

    selectors = [
        {"name": re.compile("job|posting|description|content|main|career|opening", re.I)},
        {"id": re.compile("job|posting|description|content|main|career|opening", re.I)},
        {"class": re.compile("job|posting|description|content|main|career|opening", re.I)},
    ]
    for selector in selectors:
        for node in soup.find_all(attrs=selector):
            text = node.get_text("\n", strip=True)
            if len(text) > 250:
                candidates.append(text)

    main = soup.find("main") or soup.find(attrs={"role": "main"}) or soup.body or soup
    candidates.append(main.get_text("\n", strip=True))

    cleaned_candidates = [_clean_webpage_text(c) for c in candidates if c and len(c) > 100]
    if not cleaned_candidates:
        return ""
    return max(cleaned_candidates, key=_job_text_score)


def parse_job_listing(url: str) -> str:
    if not url or not url.strip():
        raise ValueError("Please enter a valid job listing URL.")
    url = url.strip()

    candidates: List[str] = []
    if AsyncChromiumLoader is not None and Html2TextTransformer is not None:
        try:
            loader = AsyncChromiumLoader([url], user_agent=os.environ.get("USER_AGENT"))
            html_docs = loader.load()
            if html_docs:
                html = html_docs[0].page_content or ""
                json_ld = _extract_json_ld_jobposting(html)
                if len(json_ld) > 300:
                    candidates.append(json_ld)
                html2text = Html2TextTransformer()
                raw_text = html2text.transform_documents(html_docs)[0].page_content
                candidates.append(raw_text)
        except Exception:
            # Browser extraction is best-effort; requests/trafilatura fallback may still work.
            pass

    try:
        fallback = _requests_fallback(url)
        if fallback:
            candidates.append(fallback)
    except Exception as exc:
        candidates.append(f"[Requests extraction failed: {exc}]")

    cleaned_candidates = [_clean_webpage_text(c) for c in candidates if c and len(c) > 200]
    raw_text = max(cleaned_candidates, key=_job_text_score, default="")

    if len(raw_text) < 250 or _job_text_score(raw_text) < 30:
        raise ValueError("Could not extract enough job description text from this URL. Paste the description manually.")

    if job_listing_prompt is not None:
        try:
            agent = get_agent(temperature=0.1)
            prompt = job_listing_prompt.prompt_template.format(raw_text=raw_text[:18000])
            response = agent.invoke(prompt)
            cleaned = _clean_webpage_text(response.content or "")
            if len(cleaned) > 250 and _job_text_score(cleaned) >= 25:
                return cleaned
        except Exception:
            # LLM cleanup is a quality step only. Return the best deterministic extraction instead.
            pass

    return raw_text[:18000]


def format_fit_summary_for_prompt(fit_summary: Optional[Dict[str, object]]) -> str:
    if not fit_summary:
        return "No keyword fit analysis was provided. Use only the resume and job listing."

    def _items(key: str, limit: int = 12) -> str:
        values = fit_summary.get(key, [])
        if not isinstance(values, list) or not values:
            return "None detected"
        return ", ".join(str(v) for v in values[:limit])

    requirement_lines = fit_summary.get("requirement_lines", [])
    if isinstance(requirement_lines, list) and requirement_lines:
        requirements = "\n".join(f"- {line}" for line in requirement_lines[:6])
    else:
        requirements = "None detected"

    return f"""
Keyword fit score: {fit_summary.get('fit_score', 'N/A')}%
Matched resume/job skills: {_items('matched_skills')}
Job skills not clearly found in resume: {_items('missing_skills')}
Important job requirement lines to address when resume-supported:
{requirements}
""".strip()


def _generation_options_for_prompt(length: str = "Medium") -> str:
    """Translate UI choices into explicit prompt rules so output changes are visible."""
    normalized_length = (length or "Medium").strip().lower()
    length_rules = {
        "short": "Length: Short. Write 120-170 words. Use 2 compact paragraphs. Be direct and avoid extra background.",
        "medium": "Length: Medium. Write 180-250 words. Use 3 focused paragraphs with a clear opening, evidence, and closing.",
        "detailed": "Length: Detailed. Write 280-380 words. Use 4 paragraphs and include 2-3 specific resume-backed examples tied to the job requirements.",
    }
    return length_rules.get(normalized_length, length_rules["medium"])


def _friendly_generation_error(exc: Exception) -> str:
    """Convert common LLM/API failures into user-actionable messages."""
    message = str(exc) or exc.__class__.__name__
    low = message.lower()
    if "api key" in low or "authentication" in low or "unauthorized" in low:
        return "OpenAI authentication failed. Check OPENAI_API_KEY in your .env file."
    if "rate" in low or "429" in low:
        return "OpenAI rate limit was reached. Wait a moment and try again."
    if "quota" in low or "billing" in low or "insufficient" in low:
        return "OpenAI quota or billing is unavailable for this API key."
    if "timeout" in low or "connection" in low or "network" in low:
        return "Network/API connection failed. Check your internet connection and try again."
    return f"Cover letter generation failed: {message}"


def generate_cover_letter(
    resume_text: str,
    job_listing_text: str,
    prompt_template,
    fit_summary: Optional[Dict[str, object]] = None,
    length: str = "Medium",
    max_retries: int = 2,
) -> str:
    if not resume_text or not resume_text.strip():
        raise ValueError("Resume text is missing.")
    if not job_listing_text or not job_listing_text.strip():
        raise ValueError("Job listing text is missing.")
    generation_options = _generation_options_for_prompt(length)
    prompt = prompt_template.format(
        resume=resume_text[:12000],
        job_listing=job_listing_text[:12000],
        fit_summary=format_fit_summary_for_prompt(fit_summary),
        generation_options=generation_options,
    )

    last_error: Optional[Exception] = None
    for attempt in range(max(1, max_retries + 1)):
        try:
            agent = get_agent()
            response = agent.invoke(prompt)
            content = getattr(response, "content", "") or ""
            content = content.strip()
            if not content:
                raise RuntimeError("The model returned an empty response.")
            return normalize_cover_letter_closing(content)
        except Exception as exc:  # pragma: no cover - external API behavior
            last_error = exc
            low = str(exc).lower()
            retryable = any(token in low for token in ["rate", "429", "timeout", "connection", "temporarily", "network"])
            if not retryable or attempt >= max_retries:
                break
            time.sleep(1.5 * (attempt + 1))

    raise RuntimeError(_friendly_generation_error(last_error or RuntimeError("Unknown error")))


def _boundary_contains(text: str, phrase: str) -> bool:
    phrase_low = phrase.lower().strip()
    aliases = set(ALIASES.get(phrase, [])) | {phrase_low}
    canonical = _CANONICAL_BY_ALIAS.get(phrase_low)
    if canonical:
        aliases.update(ALIASES.get(canonical, []))
        aliases.add(canonical.lower())
    text_low = f" {_normalize_text(text).lower()} "
    for alias in aliases:
        alias = alias.lower().strip()
        if not alias:
            continue
        if re.search(r"(?<![a-z0-9+#.])" + re.escape(alias) + r"(?![a-z0-9+#.])", text_low):
            return True
    return False


def _canonicalize_skill(skill: str) -> str:
    skill = re.sub(r"\s+", " ", str(skill or "").strip(" ,.;:()[]{}"))
    if not skill:
        return ""
    low = skill.lower()
    return _CANONICAL_BY_ALIAS.get(low, skill)


def extract_catalog_skills(text: str) -> Dict[str, List[str]]:
    result: Dict[str, List[str]] = OrderedDict()
    for category, skills in SKILL_CATALOG.items():
        hits = [_canonicalize_skill(skill) for skill in skills if _boundary_contains(text, skill)]
        hits = sorted(set(hits), key=str.lower)
        if hits:
            result[category] = hits
    return result


def _flatten_skill_map(skill_map: Dict[str, List[str]]) -> List[str]:
    skills: List[str] = []
    for values in skill_map.values():
        skills.extend(values)
    return sorted(set(skills), key=str.lower)


_REQUIREMENT_SECTION_HEADINGS = {
    "required qualifications",
    "minimum qualifications",
    "basic qualifications",
    "preferred qualifications",
    "qualifications",
    "requirements",
    "responsibilities",
    "what you'll do",
    "what you will do",
    "what we're looking for",
    "what we are looking for",
}


def _is_requirement_heading(line: str) -> bool:
    normalized = re.sub(r"[^a-z0-9' ]+", "", (line or "").lower()).strip()
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized in _REQUIREMENT_SECTION_HEADINGS or (len(normalized.split()) <= 4 and normalized.endswith("qualifications"))


def _extract_requirement_lines(job_text: str) -> List[str]:
    lines = []
    for raw in re.split(r"[\n•·]|(?<=\.)\s+", job_text or ""):
        line = raw.strip(" -–—\t:")
        if _is_requirement_heading(line):
            continue
        if 25 <= len(line) <= 240:
            low = line.lower()
            if any(marker in low for marker in ["experience", "proficient", "knowledge", "familiar", "required", "preferred", "build", "develop", "design", "sql", "python", "api", "data", "cloud", "degree", "framework", "database"]):
                lines.append(line)
    seen, unique = set(), []
    for line in lines:
        key = line.lower()
        if key not in seen:
            seen.add(key)
            unique.append(line)
    return unique[:10]


def _extract_dynamic_job_terms(job_text: str) -> List[str]:
    """Capture important tech-looking terms not already in the catalog without adding generic noise."""
    terms: Set[str] = set()
    patterns = [
        r"\b[A-Z][A-Za-z0-9]*(?:\.[A-Za-z0-9]+)+\b",      # Node.js, Vue.js
        r"\b[A-Za-z]+(?:[+#]{1,2})\b",                       # C++, C#
        r"\b[A-Z]{2,}(?:/[A-Z]{2,})?\b",                     # AWS, CI/CD
        r"\b[A-Za-z]+-[A-Za-z0-9-]+\b",                      # scikit-learn
    ]
    for pattern in patterns:
        for match in re.findall(pattern, job_text or ""):
            term = _canonicalize_skill(match)
            low = term.lower()
            if 1 < len(term) < 35 and low not in GENERIC_STOPWORDS:
                terms.add(term)
    return sorted(terms, key=str.lower)[:16]


def _llm_extract_skills(text: str, source_name: str) -> List[str]:
    if not os.getenv("OPENAI_API_KEY"):
        return []
    try:
        agent = get_agent(temperature=0)
        prompt = f"""
Extract only concrete technical skills, tools, programming languages, frameworks, databases, cloud/devops tools, BI tools, and data/ML methods from the {source_name} text.
Return strict JSON only: {{"skills": ["Python", "SQL"]}}
Do not include soft skills, responsibilities, job benefits, company values, generic nouns, or phrases like communication/teamwork.

Text:
{text[:9000]}
"""
        response = agent.invoke(prompt)
        raw = (response.content or "").strip()
        match = re.search(r"\{.*\}", raw, flags=re.S)
        if match:
            raw = match.group(0)
        data = json.loads(raw)
        skills = data.get("skills", []) if isinstance(data, dict) else []
        clean = []
        for skill in skills:
            skill = _canonicalize_skill(skill)
            if 1 < len(skill) < 45 and skill.lower() not in GENERIC_STOPWORDS:
                clean.append(skill)
        return sorted(set(clean), key=str.lower)[:28]
    except Exception:
        return []


def _filter_job_skills(skills: Iterable[str], job_text: str) -> Set[str]:
    filtered: Set[str] = set()
    for skill in skills:
        skill = _canonicalize_skill(skill)
        if not skill or skill.lower() in GENERIC_STOPWORDS:
            continue
        if _boundary_contains(job_text, skill) or skill in _flatten_skill_map(extract_catalog_skills(job_text)):
            filtered.add(skill)
    return filtered


def analyze_resume_job_fit(resume_text: str, job_listing_text: str) -> Dict[str, object]:
    resume_catalog = extract_catalog_skills(resume_text)
    job_catalog = extract_catalog_skills(job_listing_text)

    resume_skills: Set[str] = set(_flatten_skill_map(resume_catalog))
    job_skills: Set[str] = set(_flatten_skill_map(job_catalog))

    llm_job = _filter_job_skills(_llm_extract_skills(job_listing_text, "job description"), job_listing_text)
    llm_resume = {_canonicalize_skill(s) for s in _llm_extract_skills(resume_text, "resume")}
    job_skills.update(llm_job)
    resume_skills.update(s for s in llm_resume if s and s.lower() not in GENERIC_STOPWORDS)

    # Add only tech-looking dynamic terms from the JD.
    job_skills.update(_extract_dynamic_job_terms(job_listing_text))

    matched = sorted([skill for skill in job_skills if _boundary_contains(resume_text, skill) or skill in resume_skills], key=str.lower)
    missing = sorted([skill for skill in job_skills if skill not in matched], key=str.lower)

    fit_score = round((len(matched) / len(job_skills)) * 100) if job_skills else 0
    fit_score = max(0, min(100, fit_score))

    return {
        "fit_score": fit_score,
        "matched_skills": matched[:24],
        "missing_skills": missing[:24],
        "job_skills_detected": sorted(job_skills, key=str.lower),
        "resume_skills_detected": sorted(resume_skills, key=str.lower),
        "job_skills_by_category": job_catalog,
        "resume_skills_by_category": resume_catalog,
        "requirement_lines": _extract_requirement_lines(job_listing_text),
    }


_CLOSING_START_RE = re.compile(
    r"^(sincerely|best regards|kind regards|regards|yours sincerely),?\s*(.*)$",
    re.IGNORECASE,
)


def _format_closing_block(lines: List[str]) -> str:
    """Keep the closing phrase and candidate name on separate lines for exports."""
    if not lines:
        return ""

    first = lines[0].strip()
    match = _CLOSING_START_RE.match(first)
    if not match:
        return " ".join(line.strip() for line in lines if line.strip())

    closing_phrase = match.group(1).strip().title()
    inline_name = match.group(2).strip()
    name_parts = []
    if inline_name:
        name_parts.append(inline_name)
    if len(lines) > 1:
        name_parts.extend(line.strip() for line in lines[1:] if line.strip())

    candidate_name = " ".join(name_parts).strip()
    if candidate_name:
        return f"{closing_phrase},\n{candidate_name}"
    return f"{closing_phrase},"


def _cover_letter_paragraphs(text: str) -> List[str]:
    """Return clean paragraphs while preserving the sign-off/name line break."""
    paragraphs: List[str] = []
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n").strip()

    for block in re.split(r"\n\s*\n", normalized):
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue

        closing_index = next((idx for idx, line in enumerate(lines) if _CLOSING_START_RE.match(line)), None)
        if closing_index is None:
            paragraphs.append(" ".join(lines))
            continue

        before_closing = " ".join(lines[:closing_index]).strip()
        if before_closing:
            paragraphs.append(before_closing)

        closing_block = _format_closing_block(lines[closing_index:])
        if closing_block:
            paragraphs.append(closing_block)

    if not paragraphs:
        paragraphs = [line.strip() for line in normalized.splitlines() if line.strip()]
    return paragraphs


def normalize_cover_letter_closing(text: str) -> str:
    """Normalize common inline sign-offs so exports preserve the name on a new line."""
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    # Convert: "Sincerely, Jane Doe" -> "Sincerely,\nJane Doe" near the end of the letter.
    pattern = re.compile(
        r"(?im)^(\s*(?:sincerely|best regards|kind regards|regards|yours sincerely),)\s+([^\n]+?)\s*$"
    )
    return pattern.sub(lambda m: f"{m.group(1)}\n{m.group(2).strip()}", text)


def build_export_metadata(
    candidate_name: str = "",
    email: str = "",
    phone: str = "",
    linkedin: str = "",
    portfolio: str = "",
    company: str = "",
    job_title: str = "",
    include_header: bool = True,
) -> Dict[str, object]:
    """Create sanitized optional metadata for DOCX/PDF exports."""
    def clean(value: str) -> str:
        return re.sub(r"\s+", " ", str(value or "")).strip()

    return {
        "candidate_name": clean(candidate_name),
        "email": clean(email),
        "phone": clean(phone),
        "linkedin": clean(linkedin),
        "portfolio": clean(portfolio),
        "company": clean(company),
        "job_title": clean(job_title),
        "date": date.today().strftime("%B %-d, %Y") if os.name != "nt" else date.today().strftime("%B %#d, %Y"),
        "include_header": bool(include_header),
    }


def _metadata_contact_line(metadata: Optional[Dict[str, object]]) -> str:
    if not metadata:
        return ""
    parts = [
        str(metadata.get("email") or "").strip(),
        str(metadata.get("phone") or "").strip(),
        str(metadata.get("linkedin") or "").strip(),
        str(metadata.get("portfolio") or "").strip(),
    ]
    return " | ".join(part for part in parts if part)


def _metadata_has_export_header(metadata: Optional[Dict[str, object]]) -> bool:
    if not metadata or not metadata.get("include_header"):
        return False
    return any(str(metadata.get(key) or "").strip() for key in ["candidate_name", "email", "phone", "linkedin", "portfolio", "company", "job_title"])


def _metadata_company_line(metadata: Optional[Dict[str, object]]) -> str:
    if not metadata:
        return ""
    job_title = str(metadata.get("job_title") or "").strip()
    company = str(metadata.get("company") or "").strip()
    if job_title and company:
        return f"{company} — {job_title}"
    return company or job_title


def cover_letter_to_docx(text: str, metadata: Optional[Dict[str, object]] = None) -> bytes:
    """Create a clean, recruiter-ready DOCX export."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(9)

    if _metadata_has_export_header(metadata):
        name = str(metadata.get("candidate_name") or "").strip()
        contact = _metadata_contact_line(metadata)
        company_line = _metadata_company_line(metadata)
        current_date = str(metadata.get("date") or "").strip()

        if name:
            p = doc.add_paragraph()
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(13)
            p.paragraph_format.space_after = Pt(2)
        if contact:
            p = doc.add_paragraph(contact)
            p.paragraph_format.space_after = Pt(8)
        if current_date:
            p = doc.add_paragraph(current_date)
            p.paragraph_format.space_after = Pt(8)
        if company_line:
            p = doc.add_paragraph(company_line)
            p.paragraph_format.space_after = Pt(14)

    for paragraph in _cover_letter_paragraphs(normalize_cover_letter_closing(text)):
        p = doc.add_paragraph()
        for idx, line in enumerate(paragraph.split("\n")):
            if idx:
                p.add_run().add_break()
            p.add_run(line)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.08

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def cover_letter_to_pdf(text: str, metadata: Optional[Dict[str, object]] = None) -> bytes:
    """Create a polished PDF export with cover-letter spacing and margins."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        rightMargin=0.85 * inch,
        leftMargin=0.85 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title="Cover Letter",
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "CoverLetterBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=11.2,
        leading=16,
        spaceAfter=12,
        firstLineIndent=0,
        alignment=0,
    )
    salutation = ParagraphStyle(
        "CoverLetterSalutation",
        parent=body,
        spaceAfter=14,
    )
    closing = ParagraphStyle(
        "CoverLetterClosing",
        parent=body,
        spaceBefore=8,
        spaceAfter=6,
    )
    header_name = ParagraphStyle(
        "CoverLetterHeaderName",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        spaceAfter=2,
    )
    header_line = ParagraphStyle(
        "CoverLetterHeaderLine",
        parent=body,
        fontSize=10.2,
        leading=13,
        spaceAfter=8,
    )

    story = []
    if _metadata_has_export_header(metadata):
        name = str(metadata.get("candidate_name") or "").strip()
        contact = _metadata_contact_line(metadata)
        company_line = _metadata_company_line(metadata)
        current_date = str(metadata.get("date") or "").strip()

        if name:
            story.append(Paragraph(html_escape(name), header_name))
        if contact:
            story.append(Paragraph(html_escape(contact), header_line))
        if current_date:
            story.append(Paragraph(html_escape(current_date), body))
        if company_line:
            story.append(Paragraph(html_escape(company_line), body))
        story.append(Spacer(1, 8))

    paragraphs = _cover_letter_paragraphs(normalize_cover_letter_closing(text))
    for idx, para in enumerate(paragraphs):
        safe_para = html_escape(para).replace("\n", "<br/>")
        low = para.lower().strip()
        if idx == 0 and low.startswith(("dear ", "hello ", "hi ")):
            style = salutation
        elif low.startswith(("best regards", "sincerely", "thank you", "regards")):
            style = closing
        else:
            style = body
        story.append(Paragraph(safe_para, style))
        if idx < len(paragraphs) - 1:
            story.append(Spacer(1, 2))

    if not story:
        story.append(Paragraph("Cover letter text was empty.", body))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()
