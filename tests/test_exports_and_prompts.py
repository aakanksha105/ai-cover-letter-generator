from prompts.cover_letter_prompts import prompt_template_classic
import io

from docx import Document
from pypdf import PdfReader

from utils.utils import _cover_letter_paragraphs, cover_letter_to_docx, cover_letter_to_pdf


def test_prompt_accepts_fit_summary_and_generation_options():
    rendered = prompt_template_classic.format(
        resume="Python developer",
        job_listing="Build APIs",
        fit_summary="Matched: Python",
        generation_options="Length: Short. Write 120-170 words.",
    )
    assert "Matched: Python" in rendered
    assert "Length: Short" in rendered


def test_docx_export_returns_bytes():
    data = cover_letter_to_docx("Hello\nWorld")
    assert isinstance(data, bytes)
    assert data[:2] == b"PK"


def test_pdf_export_returns_bytes():
    data = cover_letter_to_pdf("Hello World")
    assert isinstance(data, bytes)
    assert data.startswith(b"%PDF")


def test_cover_letter_paragraphs_preserve_closing_name_line_break():
    paragraphs = _cover_letter_paragraphs("Dear Hiring Manager,\n\nThank you.\nSincerely,\nAakanksha Prashant Bhondve")
    assert paragraphs[-1] == "Sincerely,\nAakanksha Prashant Bhondve"

    inline_paragraphs = _cover_letter_paragraphs("Dear Hiring Manager,\n\nThank you.\nSincerely, Aakanksha Prashant Bhondve")
    assert inline_paragraphs[-1] == "Sincerely,\nAakanksha Prashant Bhondve"


def test_docx_export_keeps_name_after_sincerely_on_new_line():
    data = cover_letter_to_docx("Dear Hiring Manager,\n\nThank you.\nSincerely,\nAakanksha Prashant Bhondve")
    doc = Document(io.BytesIO(data))
    assert doc.paragraphs[-1].text == "Sincerely,\nAakanksha Prashant Bhondve"


def test_pdf_export_keeps_name_after_sincerely_on_new_line():
    data = cover_letter_to_pdf("Dear Hiring Manager,\n\nThank you.\nSincerely,\nAakanksha Prashant Bhondve")
    reader = PdfReader(io.BytesIO(data))
    extracted = reader.pages[0].extract_text()
    assert "Sincerely,\nAakanksha Prashant Bhondve" in extracted

from utils.utils import build_export_metadata, normalize_cover_letter_closing


def test_normalize_cover_letter_closing_handles_inline_signature():
    text = normalize_cover_letter_closing("Dear Hiring Manager,\n\nThank you.\n\nSincerely, Jane Doe")
    assert "Sincerely,\nJane Doe" in text


def test_docx_export_can_include_professional_header():
    metadata = build_export_metadata(
        candidate_name="Jane Doe",
        email="jane@example.com",
        linkedin="linkedin.com/in/janedoe",
        company="Acme",
        job_title="Software Engineer",
        include_header=True,
    )
    data = cover_letter_to_docx("Dear Hiring Manager,\n\nThank you.\n\nSincerely,\nJane Doe", metadata=metadata)
    doc = Document(io.BytesIO(data))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Jane Doe" in all_text
    assert "jane@example.com | linkedin.com/in/janedoe" in all_text
    assert "Acme — Software Engineer" in all_text


def test_pdf_export_can_include_professional_header():
    metadata = build_export_metadata(candidate_name="Jane Doe", email="jane@example.com", company="Acme")
    data = cover_letter_to_pdf("Dear Hiring Manager,\n\nThank you.\n\nSincerely,\nJane Doe", metadata=metadata)
    reader = PdfReader(io.BytesIO(data))
    extracted = reader.pages[0].extract_text()
    assert "Jane Doe" in extracted
    assert "jane@example.com" in extracted
    assert "Acme" in extracted


def test_export_header_is_skipped_when_include_header_false():
    metadata = build_export_metadata(
        candidate_name="Jane Doe",
        email="jane@example.com",
        company="Acme",
        include_header=False,
    )
    data = cover_letter_to_docx("Dear Hiring Manager,\n\nThank you.\n\nSincerely,\nJane Doe", metadata=metadata)
    doc = Document(io.BytesIO(data))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "jane@example.com" not in all_text
    assert "Acme" not in all_text
