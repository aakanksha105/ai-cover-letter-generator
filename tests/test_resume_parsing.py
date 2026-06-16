from io import BytesIO
from docx import Document

from utils.utils import _read_docx, assess_resume_extraction_quality, normalize_pasted_resume_text


def test_docx_parser_reads_table_cells():
    doc = Document()
    doc.add_paragraph("Aakanksha Example")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python, SQL, Flask"
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    text = _read_docx(buf)

    assert "Aakanksha Example" in text
    assert "Python, SQL, Flask" in text


def test_resume_quality_flags_short_text():
    quality = assess_resume_extraction_quality("Python SQL")
    assert quality["is_likely_good"] is False
    assert quality["warnings"]


def test_normalize_pasted_resume_text_cleans_spacing():
    text = normalize_pasted_resume_text("Name   Example\r\n\r\n\r\nSkills   Python")
    assert "Name Example" in text
    assert "\n\n\n" not in text


def test_resume_quality_flags_unusual_characters():
    text = "Experience Skills Education Projects " + ("�" * 80)
    quality = assess_resume_extraction_quality(text)
    assert quality["is_likely_good"] is False
    assert any("unusual characters" in warning for warning in quality["warnings"])
